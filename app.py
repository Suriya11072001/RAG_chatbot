import streamlit as st
from streamlit.runtime.uploaded_file_manager import UploadedFile
import faiss
import os
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.chains import ConversationalRetrievalChain
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_huggingface import HuggingFaceEndpoint, ChatHuggingFace
from langchain.memory import ConversationBufferMemory
from secret_api_keys import huggingface_api_key  # your HF key

# set Hugging Face API key
os.environ['HUGGINGFACEHUB_API_TOKEN'] = huggingface_api_key


def process_input(input_type, input_data):
    """Processes different input types and returns a vectorstore."""

    if input_type == "Link":
        documents = []
        for url in input_data:
            loader = WebBaseLoader(url)
            documents.extend(loader.load())
        text = "\n".join([doc.page_content for doc in documents])

    elif input_type == "PDF":
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for PDF")

        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()

    elif input_type == "Text":
        if isinstance(input_data, str):
            text = input_data
        else:
            raise ValueError("Expected a string for 'Text' input type.")

    elif input_type == "DOCX":
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        elif isinstance(input_data, UploadedFile):
            doc = Document(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for DOCX")
        text = "\n".join([para.text for para in doc.paragraphs])

    elif input_type == "TXT":
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        elif isinstance(input_data, UploadedFile):
            text = input_data.read().decode('utf-8')
        else:
            raise ValueError("Invalid input data for TXT")

    else:
        raise ValueError("Unsupported input type")

    # split into chunks
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    texts = text_splitter.split_text(text)

    # embeddings
    model_name = "sentence-transformers/all-mpnet-base-v2"
    hf_embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': False}
    )

    # create FAISS index
    sample_embedding = np.array(hf_embeddings.embed_query("sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)

    vector_store = FAISS(
        embedding_function=hf_embeddings,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
    )
    vector_store.add_texts(texts)

    return vector_store


def create_conversational_chain(vectorstore):
    """Creates a Conversational Retrieval QA chain with memory."""

    # First, create endpoint (completion model)
    llm_endpoint = HuggingFaceEndpoint(
        repo_id="meta-llama/Llama-3.1-8B-Instruct",
        task="text-generation",
        temperature=0.6,
        max_new_tokens=512,
    )

    # Wrap endpoint in ChatHuggingFace
    llm = ChatHuggingFace(llm=llm_endpoint, verbose=True)

    # Memory for chat history
    memory = ConversationBufferMemory(
        memory_key="chat_history",
        return_messages=True
    )

    # Conversational chain
    qa_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(),
        memory=memory
    )
    return qa_chain


def main():
    st.title("🧠 Conversational RAG Q&A App")

    input_type = st.selectbox("Input Type", ["Link", "PDF", "Text", "DOCX", "TXT"])

    if input_type == "Link":
        number_input = st.number_input("Enter the number of Links", min_value=1, max_value=20, step=1)
        input_data = []
        for i in range(number_input):
            url = st.text_input(f"URL {i+1}")
            if url:
                input_data.append(url)

    elif input_type == "Text":
        input_data = st.text_area("Enter the text")

    elif input_type == 'PDF':
        input_data = st.file_uploader("Upload a PDF file", type=["pdf"])

    elif input_type == 'TXT':
        input_data = st.file_uploader("Upload a text file", type=['txt'])

    elif input_type == 'DOCX':
        input_data = st.file_uploader("Upload a DOCX file", type=['docx', 'doc'])

    if st.button("Proceed") and input_data:
        vectorstore = process_input(input_type, input_data)
        st.session_state["qa_chain"] = create_conversational_chain(vectorstore)
        st.success("✅ Data processed and conversational chain created!")

    if "qa_chain" in st.session_state:
        query = st.text_input("💬 Ask your question")
        if st.button("Submit") and query:
            response = st.session_state["qa_chain"]({"question": query})
            st.write("**Answer:**", response["answer"])

            # Display chat history
            with st.expander("Chat History"):
                for i, msg in enumerate(response["chat_history"]):
                    role = "🧑 You" if i % 2 == 0 else "🤖 Assistant"
                    st.markdown(f"**{role}:** {msg.content}")


if __name__ == "__main__":
    main()
